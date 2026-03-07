"""
Document Parsing Service — Extract text from uploaded files.

Supports:
  - PDF via pymupdf4llm (markdown-structured output)
  - DOCX via python-docx (paragraphs + tables)
  - TXT / Markdown (plain UTF-8)
  - Images via Gemini Flash (vision OCR)

Runs in BackgroundTask after upload confirmation.
Updates spark_uploads.status: uploaded → parsing → parsed | failed.
"""

from __future__ import annotations

import io
import logging
from uuid import UUID

from app.config import settings
from app.services.supabase import get_supabase_client

logger = logging.getLogger(__name__)

# Max pages for PDF extraction (safety cap)
_MAX_PDF_PAGES = 200


async def parse_upload(upload_id: UUID, client_id: UUID) -> None:
    """Parse an uploaded file and store extracted text.

    Fetches the file from Supabase Storage, extracts text based on
    mime_type, and updates the spark_uploads row with parsed_text.
    """
    sb = await get_supabase_client()

    # Mark as parsing
    await (
        sb.table("spark_uploads")
        .update({"status": "parsing"})
        .eq("id", str(upload_id))
        .eq("client_id", str(client_id))
        .execute()
    )

    try:
        # Fetch upload row
        result = await (
            sb.table("spark_uploads")
            .select("*")
            .eq("id", str(upload_id))
            .eq("client_id", str(client_id))
            .limit(1)
            .execute()
        )

        if not result.data:
            logger.error("Upload %s not found during parsing", upload_id)
            return

        row = result.data[0]
        mime_type: str = row["mime_type"]
        storage_path: str = row["storage_path"]

        # Download file from storage
        file_bytes = await sb.storage.from_(
            settings.supabase_storage_bucket
        ).download(storage_path)

        # Route by mime type
        if mime_type == "application/pdf":
            text, page_count = _parse_pdf(file_bytes)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text, page_count = _parse_docx(file_bytes)
        elif mime_type in ("text/plain", "text/markdown"):
            text, page_count = _parse_txt(file_bytes)
        elif mime_type.startswith("image/"):
            text, page_count = await _parse_image(file_bytes, mime_type)
        else:
            raise ValueError(f"Unsupported mime type: {mime_type}")

        # Update row with parsed text
        await (
            sb.table("spark_uploads")
            .update(
                {
                    "status": "parsed",
                    "parsed_text": text,
                    "page_count": page_count,
                }
            )
            .eq("id", str(upload_id))
            .eq("client_id", str(client_id))
            .execute()
        )

        logger.info(
            "Parsed upload %s: %d chars, %s pages",
            upload_id,
            len(text),
            page_count,
        )

    except Exception as e:
        logger.exception("Failed to parse upload %s", upload_id)
        await (
            sb.table("spark_uploads")
            .update(
                {
                    "status": "failed",
                    "error_message": str(e)[:500],
                }
            )
            .eq("id", str(upload_id))
            .eq("client_id", str(client_id))
            .execute()
        )


def _parse_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Extract markdown-structured text from a PDF using pymupdf4llm."""
    import pymupdf4llm
    import pymupdf

    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    page_count = min(len(doc), _MAX_PDF_PAGES)
    doc.close()

    # pymupdf4llm produces clean markdown from PDFs
    text = pymupdf4llm.to_markdown(
        pymupdf.open(stream=file_bytes, filetype="pdf"),
        pages=list(range(page_count)),
    )

    if not text or not text.strip():
        # Likely a scanned PDF — fall back to page text extraction
        doc = pymupdf.open(stream=file_bytes, filetype="pdf")
        parts = []
        for i in range(page_count):
            page_text = doc[i].get_text()
            if page_text.strip():
                parts.append(page_text)
        doc.close()
        text = "\n\n".join(parts)

    return text.strip(), page_count


def _parse_docx(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a DOCX file preserving structure."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Preserve heading structure with markdown-style prefixes
        style_name = (para.style.name or "").lower()
        if "heading 1" in style_name:
            parts.append(f"# {text}")
        elif "heading 2" in style_name:
            parts.append(f"## {text}")
        elif "heading 3" in style_name:
            parts.append(f"### {text}")
        elif "list" in style_name:
            parts.append(f"- {text}")
        else:
            parts.append(text)

    # Extract tables
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    # DOCX doesn't have a native page count — estimate from paragraphs
    estimated_pages = max(1, len(parts) // 30)

    return "\n\n".join(parts).strip(), estimated_pages


def _parse_txt(file_bytes: bytes) -> tuple[str, int]:
    """Extract text from a plain text file."""
    # Handle BOM
    text = file_bytes.decode("utf-8-sig").strip()
    estimated_pages = max(1, len(text) // 3000)
    return text, estimated_pages


async def _parse_image(file_bytes: bytes, mime_type: str) -> tuple[str, int]:
    """Extract text from an image using Gemini Flash vision."""
    import base64

    from app.services import llm

    b64 = base64.b64encode(file_bytes).decode("ascii")

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{mime_type};base64,{b64}",
                    },
                },
                {
                    "type": "text",
                    "text": (
                        "Extract all text content from this image. "
                        "If it contains a document, preserve the structure "
                        "(headings, lists, paragraphs). If it's a photo or "
                        "graphic, describe the visual content and any visible text. "
                        "Return the extracted content only, no commentary."
                    ),
                },
            ],
        }
    ]

    text = await llm.complete(
        messages=messages,
        model=settings.spark_primary_model,  # Gemini Flash — cheap vision
        temperature=0.1,
        timeout=30,
    )

    return text.strip(), 1
